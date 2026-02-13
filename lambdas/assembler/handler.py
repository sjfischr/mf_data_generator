"""Document Assembler Lambda — combines sections into final DOCX package."""

import io
import json
import logging
import os
import re
import tempfile
import zipfile

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from lambdas.shared import s3_utils
from lambdas.shared.models import CrosswalkData

logger = logging.getLogger(__name__)

SECTION_ORDER = [
    "section_01_introduction",
    "section_02_property_description",
    "section_03_market_analysis",
    "section_04_highest_best_use",
    "section_05_valuation_methodology",
    "section_06_sales_comparison",
    "section_07_income_approach",
    "section_08_cost_approach",
    "section_09_reconciliation",
    "section_10_assumptions",
    "section_11_certification",
    "section_12_addenda",
]

LENDER_NAME = "NorthBridge Multifamily Capital"
REPORT_TITLE = "Synthetic Multifamily Appraisal Report"
STOPWORDS = {
    "a", "an", "and", "at", "by", "for", "from", "in", "into", "is", "of", "on", "or", "the", "to", "with",
}


def get_cover_metadata(job_id: str) -> tuple[str, str]:
    """Resolve property name and location text for the cover page."""
    property_name = "Subject Property"
    location_text = ""

    try:
        crosswalk_data = s3_utils.read_json(job_id, "crosswalk-data.json")
        crosswalk = CrosswalkData.model_validate(crosswalk_data)
        prop = crosswalk.property_identification
        property_name = prop.property_name or property_name
        location_text = ", ".join(
            [part for part in [prop.address, prop.city, prop.state] if part]
        )
    except Exception:
        try:
            input_data = s3_utils.read_json(job_id, "input.json")
            property_name = input_data.get("property_name") or property_name
            location_text = ", ".join(
                [
                    part
                    for part in [
                        input_data.get("address", ""),
                        input_data.get("city", ""),
                        input_data.get("state", ""),
                    ]
                    if part
                ]
            )
        except Exception:
            pass

    return property_name, location_text


def select_cover_image_filename(manifest: list[dict]) -> str | None:
    """Choose the best generated image for the report cover."""
    successful = [entry for entry in manifest if entry.get("status") == "success"]
    if not successful:
        return None

    preferred_keywords = ("aerial", "exterior", "front", "facade", "building")
    for entry in successful:
        haystack = f"{entry.get('filename', '')} {entry.get('description', '')}".lower()
        if any(keyword in haystack for keyword in preferred_keywords):
            return entry.get("filename")

    return successful[0].get("filename")


def add_cover_page(doc: Document, property_name: str, location_text: str, cover_image_path: str | None):
    """Insert a branded cover page at the beginning of the report."""
    lender = doc.add_paragraph()
    lender.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lender_run = lender.add_run(LENDER_NAME)
    lender_run.bold = True
    lender_run.font.size = Pt(22)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(REPORT_TITLE)
    title_run.bold = True
    title_run.font.size = Pt(16)

    doc.add_paragraph()

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(property_name)
    name_run.bold = True
    name_run.font.size = Pt(28)

    if location_text:
        location_para = doc.add_paragraph()
        location_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        location_run = location_para.add_run(location_text)
        location_run.font.size = Pt(12)

    doc.add_paragraph()

    if cover_image_path:
        try:
            doc.add_picture(cover_image_path, width=Inches(6.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as exc:
            logger.warning("Could not add cover image: %s", exc)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run("Prepared for synthetic loan package review")
    date_run.italic = True
    date_run.font.size = Pt(10)

    doc.add_page_break()


def collect_sections(job_id: str) -> str:
    """Read all section markdown files and combine in order."""
    combined = []

    for section_name in SECTION_ORDER:
        filename = f"sections/{section_name}.md"
        try:
            content = s3_utils.read_text(job_id, filename)
            combined.append(content)
            combined.append("\n\n\\newpage\n\n")
        except Exception as e:
            logger.warning("Section %s not found: %s", section_name, e)
            combined.append(f"\n\n*[Section {section_name} not available]*\n\n")

    return "\n".join(combined)


def insert_images(markdown: str, job_id: str) -> str:
    """Replace [IMAGE: description] placeholders with actual image references."""
    try:
        manifest = s3_utils.read_json(job_id, "images/manifest.json")
    except Exception:
        logger.warning("Image manifest not found, skipping image insertion")
        return markdown

    image_entries = [entry for entry in manifest if entry.get("status") == "success" and entry.get("filename")]
    if not image_entries:
        logger.warning("No successful images found in manifest, skipping image insertion")
        return markdown

    used_filenames: set[str] = set()

    def tokenize(text: str) -> set[str]:
        return {
            token
            for token in re.findall(r"[a-z0-9]+", (text or "").lower())
            if token and len(token) > 2 and token not in STOPWORDS
        }

    def score_match(desc_tokens: set[str], entry: dict) -> int:
        haystack = f"{entry.get('description', '')} {entry.get('filename', '')}".lower()
        haystack_tokens = tokenize(haystack)
        overlap = len(desc_tokens & haystack_tokens)
        bonus = 0
        if "aerial" in desc_tokens and "aerial" in haystack_tokens:
            bonus += 2
        if "exterior" in desc_tokens and "exterior" in haystack_tokens:
            bonus += 2
        if "interior" in desc_tokens and "interior" in haystack_tokens:
            bonus += 2
        return overlap + bonus

    def replace_image(match):
        desc = match.group(1).strip()
        desc_tokens = tokenize(desc)

        available = [entry for entry in image_entries if entry["filename"] not in used_filenames]
        if not available:
            available = image_entries

        best_entry = max(
            available,
            key=lambda entry: (score_match(desc_tokens, entry), entry.get("filename", "")),
        )

        filename = best_entry.get("filename")
        if filename:
            used_filenames.add(filename)
            return f"![{desc}](images/{filename})"

        return f"*[Photo: {desc}]*"

    return re.sub(r"\[IMAGE:\s*([^\]]+)\]", replace_image, markdown)


def markdown_to_docx(markdown: str, job_id: str) -> bytes:
    """Convert markdown to DOCX using python-docx."""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    property_name, location_text = get_cover_metadata(job_id)

    # Download images to temp dir if needed
    image_files = {}
    cover_image_path = None
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            manifest = s3_utils.read_json(job_id, "images/manifest.json")
            img_dir = os.path.join(tmpdir, "images")
            os.makedirs(img_dir, exist_ok=True)
            
            s3 = s3_utils.get_s3_client()
            for entry in manifest:
                if entry.get("status") == "success":
                    s3_key = entry["s3_key"]
                    local_path = os.path.join(img_dir, entry["filename"])
                    s3.download_file(s3_utils.BUCKET, s3_key, local_path)
                    image_files[entry["filename"]] = local_path

            cover_image_filename = select_cover_image_filename(manifest)
            if cover_image_filename and cover_image_filename in image_files:
                cover_image_path = image_files[cover_image_filename]
        except Exception as e:
            logger.warning("Could not download images: %s", e)

        add_cover_page(doc, property_name, location_text, cover_image_path)
        
        # Parse markdown line by line
        lines = markdown.split('\n')
        i = 0
        in_list = False
        
        while i < len(lines):
            line = lines[i]
            stripped_line = line.strip()

            # Markdown tables
            if is_markdown_table_line(line):
                table_lines = []
                while i < len(lines) and is_markdown_table_line(lines[i]):
                    table_lines.append(lines[i])
                    i += 1

                if render_markdown_table(doc, table_lines):
                    in_list = False
                    continue

                # Fallback to plain text if table parsing fails
                for raw in table_lines:
                    raw_stripped = raw.strip()
                    if raw_stripped:
                        p = doc.add_paragraph()
                        add_formatted_text(p, apply_inline_formatting(raw_stripped))
                in_list = False
                continue

            # Ignore pure pipe/separator artifacts like "|||||"
            if re.match(r'^\s*\|[\s\|\-:]*\|\s*$', stripped_line):
                i += 1
                continue
            
            # Page break
            if stripped_line == '\\newpage':
                doc.add_page_break()
                i += 1
                continue
            
            # Heading 1
            if line.startswith('# '):
                text = line[2:].strip()
                doc.add_heading(text, level=1)
                i += 1
                continue
            
            # Heading 2
            if line.startswith('## '):
                text = line[3:].strip()
                doc.add_heading(text, level=2)
                i += 1
                continue
            
            # Heading 3
            if line.startswith('### '):
                text = line[4:].strip()
                doc.add_heading(text, level=3)
                i += 1
                continue
            
            # Image reference
            img_match = re.match(r'!\[([^\]]*)\]\(([^\)]+)\)', stripped_line)
            if img_match:
                alt_text = img_match.group(1)
                img_path = img_match.group(2)
                filename = os.path.basename(img_path)
                
                if filename in image_files:
                    try:
                        doc.add_picture(image_files[filename], width=Inches(5))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if alt_text:
                            caption = doc.add_paragraph(alt_text)
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption.runs[0].font.size = Pt(10)
                            caption.runs[0].font.italic = True
                    except Exception as e:
                        logger.warning("Could not insert image %s: %s", filename, e)
                        doc.add_paragraph(f"[Image: {alt_text}]")
                else:
                    doc.add_paragraph(f"[Image: {alt_text}]")
                i += 1
                continue
            
            # Bullet list item
            if stripped_line.startswith('- ') or stripped_line.startswith('* '):
                text = stripped_line[2:].strip()
                text = apply_inline_formatting(text)
                p = doc.add_paragraph(style='List Bullet')
                add_formatted_text(p, text)
                in_list = True
                i += 1
                continue
            
            # Numbered list item
            num_match = re.match(r'^\s*\d+\.\s+(.+)$', line)
            if num_match:
                text = num_match.group(1).strip()
                text = apply_inline_formatting(text)
                p = doc.add_paragraph(style='List Number')
                add_formatted_text(p, text)
                in_list = True
                i += 1
                continue
            
            # Empty line
            if not stripped_line:
                if in_list:
                    in_list = False
                i += 1
                continue
            
            # Regular paragraph
            if stripped_line:
                in_list = False
                text = apply_inline_formatting(stripped_line)
                p = doc.add_paragraph()
                add_formatted_text(p, text)
            
            i += 1
    
    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def apply_inline_formatting(text: str) -> str:
    """Preserve formatting markersto process later."""
    return text


def is_markdown_table_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if not (stripped.startswith("|") and stripped.endswith("|")):
        return False
    return stripped.count("|") >= 3


def is_markdown_separator_row(line: str) -> bool:
    stripped = line.strip()
    if not is_markdown_table_line(stripped):
        return False
    cells = [cell.strip() for cell in stripped.strip("|").split("|")]
    return all(c and re.fullmatch(r":?-{3,}:?", c) for c in cells)


def parse_markdown_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def render_markdown_table(doc: Document, table_lines: list[str]) -> bool:
    if not table_lines:
        return False

    rows = [parse_markdown_table_row(line) for line in table_lines if not is_markdown_separator_row(line)]
    rows = [row for row in rows if row]

    if len(rows) < 2:
        return False

    col_count = max(len(row) for row in rows)
    if col_count == 0:
        return False

    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for row_index, row in enumerate(rows):
        for col_index in range(col_count):
            cell_text = row[col_index] if col_index < len(row) else ""
            paragraph = table.cell(row_index, col_index).paragraphs[0]
            paragraph.clear()
            add_formatted_text(paragraph, apply_inline_formatting(cell_text))
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True

    return True


def add_formatted_text(paragraph, text: str):
    """Add text with inline formatting (bold, italic) to a paragraph."""
    # Simple regex-based parser for **bold** and *italic*
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)')
    parts = pattern.findall(text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            # Italic  
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def create_zip_package(job_id: str, docx_bytes: bytes) -> bytes:
    """Create a ZIP file containing all deliverables."""
    buf = io.BytesIO()

    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        # Add appraisal report
        zf.writestr("appraisal_report.docx", docx_bytes)

        # Add Excel files
        s3 = s3_utils.get_s3_client()
        excel_files = [
            "outputs/rent_roll.xlsx",
            "outputs/t12_year1.xlsx",
            "outputs/t12_year2.xlsx",
            "outputs/t12_year3.xlsx",
        ]
        for excel_file in excel_files:
            try:
                key = s3_utils.job_key(job_id, excel_file)
                obj = s3.get_object(Bucket=s3_utils.BUCKET, Key=key)
                filename = excel_file.split("/")[-1]
                zf.writestr(filename, obj["Body"].read())
            except Exception as e:
                logger.warning("Could not add %s to ZIP: %s", excel_file, e)

    return buf.getvalue()


def handler(event, context):
    job_id = event["job_id"]
    logger.info("Assembling final documents for job %s", job_id)

    # Collect and combine sections
    combined_markdown = collect_sections(job_id)

    # Insert images
    combined_markdown = insert_images(combined_markdown, job_id)

    # Convert to DOCX
    docx_bytes = markdown_to_docx(combined_markdown, job_id)

    # Upload DOCX
    s3_utils.write_bytes(
        job_id, "outputs/appraisal_report.docx", docx_bytes,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    # Create ZIP package
    zip_bytes = create_zip_package(job_id, docx_bytes)
    s3_utils.write_bytes(
        job_id, "outputs/loan_package.zip", zip_bytes,
        "application/zip",
    )

    # Generate presigned URLs (7 days)
    urls = {
        "appraisal": s3_utils.generate_presigned_url(job_id, "outputs/appraisal_report.docx"),
        "rent_roll": s3_utils.generate_presigned_url(job_id, "outputs/rent_roll.xlsx"),
        "t12_files": [
            s3_utils.generate_presigned_url(job_id, "outputs/t12_year1.xlsx"),
            s3_utils.generate_presigned_url(job_id, "outputs/t12_year2.xlsx"),
            s3_utils.generate_presigned_url(job_id, "outputs/t12_year3.xlsx"),
        ],
        "complete_package": s3_utils.generate_presigned_url(job_id, "outputs/loan_package.zip"),
    }

    # Save URLs to S3 for the download endpoint
    s3_utils.write_json(job_id, "outputs/download_urls.json", urls)

    logger.info("Assembly complete for job %s", job_id)

    return {
        "status": "success",
        "job_id": job_id,
        "outputs": urls,
    }
